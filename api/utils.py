from django.contrib.auth import authenticate
import json
import jwt
import requests
import docx

def jwt_get_username_from_payload_handler(payload):
    username = payload.get('sub').replace('|', '.')
    authenticate(remote_user=username)
    return username


def jwt_decode_token(token):
    header = jwt.get_unverified_header(token)
    jwks = requests.get('https://{}/.well-known/jwks.json'.format('niccko.eu.auth0.com')).json()
    public_key = None
    for jwk in jwks['keys']:
        if jwk['kid'] == header['kid']:
            public_key = jwt.algorithms.RSAAlgorithm.from_jwk(json.dumps(jwk))

    if public_key is None:
        raise Exception('Public key not found.')

    issuer = 'https://{}/'.format('niccko.eu.auth0.com')
    return jwt.decode(token, public_key, audience='https://scholarship/api/sport', issuer=issuer, algorithms=['RS256'])

def create_table(data):
    document = docx.Document()
    document.add_paragraph('Получение в течение 1-ого года, предшествующего назначению повышенной государственной академической стипендии, награды\
                            (приза) за результаты спортивной деятельности, осуществленной в рамках спортивных международных, всероссийских, ведомственных,\
                            региональных мероприятий, проводимых Университетом или иной организацией')
    object_data = json.loads(data)
    table1 = document.add_table(rows=1,cols=6)
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = 'Название мероприятия'
    hdr_cells[1].text = 'Уровень(международное,всероссийское,ведомственное и региональное)'
    hdr_cells[2].text = 'Степень участия(индивидуальное, командное)'
    hdr_cells[3].text = 'Занятое место'
    hdr_cells[4].text = 'Дата мероприятия'
    hdr_cells[5].text = 'Баллы'
    for i in object_data:
        print(i['place'])
        row_cells = table1.add_row().cells
        row_cells[0].text = i['name']
        row_cells[1].text = i['level']
        row_cells[2].text = i['degree']
        row_cells[3].text = str(i['place'])
        row_cells[4].text = i['date']
        row_cells[5].text = str(i['points'])
    document.save('Global_event.docx')


